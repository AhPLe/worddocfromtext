"""Hollow the hand-formatted Word master into template.docx (M2).

Re-runnable: if Al re-uploads an updated Word master, running this again
re-derives the template. Formatting decisions applied here (Jul 2026,
Al's picks):
  - 0.5" symmetric side margins; every right tab stop shifts by the same
    delta so date/GPA columns follow the new text width.
  - The VML autoformat horizontal-rule paragraphs are deleted; section
    headings get a bottom border instead (+6pt space-before for
    section/entry hierarchy).
  - Bullet spacing moves to the List Paragraph STYLE (6pt after — the
    current airy look) and direct per-bullet overrides are stripped, so
    spacing is a one-knob edit. Tight-groups variant = tick "Don't add
    space between paragraphs of the same style" on that style in Word.
  - Run-level font family overrides are stripped (kills Times New Roman
    paste residue); the theme (Aptos) owns the font — change it once in
    Word via Design > Fonts or the Normal style.
  - Header (name + contact line, incl. live hyperlinks) stays STATIC —
    it never varies per cut, and this preserves the hand-made links.

Content is replaced with docxtpl Jinja tags fed by the M1 render
context. One prototype block per repeating structure carries all the
formatting; loops stamp it per item.

Usage: python tools/hollow_template.py <master.docx> <template.docx>
"""

import copy
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

SECTION_HEADINGS = ("Summary", "Professional Experience", "Education",
                    "Project Experience", "Technical Skills",
                    "Publications and Awards")

MERGE_RUNS = "/mnt/skills/public/docx/scripts/merge_runs.py"


# ---------- small XML helpers ----------

def tag_paragraph(tag):
    return parse_xml(
        f'<w:p {nsdecls("w")}><w:r><w:t xml:space="preserve">{tag}</w:t>'
        '</w:r></w:p>')


def make_run(rpr_src, text=None, tab=False):
    r = parse_xml(f'<w:r {nsdecls("w")}/>')
    if rpr_src is not None:
        r.append(copy.deepcopy(rpr_src))
    if tab:
        r.append(parse_xml(f'<w:tab {nsdecls("w")}/>'))
    if text is not None:
        t = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve"/>')
        t.text = text
        r.append(t)
    return r


def rebuild_runs(p, runs):
    """Replace a paragraph's runs, keeping its pPr (formatting)."""
    for r in p._p.findall(qn("w:r")):
        p._p.remove(r)
    for r in runs:
        p._p.append(r)


def rpr(p, pred=None):
    """rPr of the first run (optionally the first matching pred)."""
    for run in p.runs:
        el = run._r.find(qn("w:rPr"))
        if el is None:
            continue
        if pred is None or pred(run):
            return el
    return None


def delete(p):
    p._p.getparent().remove(p._p)


def insert_before(anchor_p, node):
    anchor_p._p.addprevious(node)


def insert_after(anchor_p, node):
    anchor_p._p.addnext(node)


def add_bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn("w:pBdr")) is not None:
        return
    bdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="6" w:space="2" w:color="auto"/>'
        '</w:pBdr>')
    # OOXML pPr child order: pBdr precedes tabs/spacing/ind/rPr.
    for tagname in ("w:tabs", "w:spacing", "w:ind", "w:rPr"):
        anchor = pPr.find(qn(tagname))
        if anchor is not None:
            anchor.addprevious(bdr)
            return
    pPr.append(bdr)


# ---------- document lookups ----------

def find_one(doc, startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise SystemExit(f"marker not found: {startswith!r}")


def paragraphs_between(doc, start_p, end_p):
    """Paragraph objects strictly after start_p and before end_p.

    doc.paragraphs makes fresh wrapper objects on every call, so
    membership must be tested on the underlying XML elements.
    """
    ps = doc.paragraphs
    els = [p._p for p in ps]
    i, j = els.index(start_p._p), els.index(end_p._p)
    return ps[i + 1:j]


def bullets_after(doc, p):
    """Contiguous List-style paragraphs immediately following p."""
    ps = doc.paragraphs
    els = [q._p for q in ps]
    out = []
    for q in ps[els.index(p._p) + 1:]:
        if q.style.name.startswith("List"):
            out.append(q)
        else:
            break
    return out


# ---------- the surgery ----------

def hollow(src: Path, dst: Path):
    # merge fragmented runs first so text matching is reliable
    with tempfile.TemporaryDirectory() as td:
        merged = Path(td) / "merged.docx"
        subprocess.run([sys.executable, MERGE_RUNS, str(src),
                        "-o", str(merged)], check=True,
                       capture_output=True)
        doc = Document(str(merged))

        # 1) margins: 0.5" symmetric; shift every right tab by the delta
        sec = doc.sections[0]
        old_width = (sec.page_width.twips - sec.left_margin.twips
                     - sec.right_margin.twips)
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)
        new_width = (sec.page_width.twips - sec.left_margin.twips
                     - sec.right_margin.twips)
        delta = new_width - old_width
        if delta:
            for p in doc.paragraphs:
                pPr = p._p.pPr
                tabs = pPr.find(qn("w:tabs")) if pPr is not None else None
                if tabs is None:
                    continue
                for t in tabs.findall(qn("w:tab")):
                    if t.get(qn("w:val")) == "right":
                        t.set(qn("w:pos"),
                              str(int(t.get(qn("w:pos"))) + delta))

        # Strip the legacy NEGATIVE left indents that visually compensated
        # for the old 1.0" margin — once margins are normalized they push
        # content into the margin (the page-edge bug Al hit in Word,
        # Jul 2026). Bullet rows get Al's hand-fixed profile instead.
        for p in doc.paragraphs:
            pPr = p._p.pPr
            ind = pPr.find(qn("w:ind")) if pPr is not None else None
            if ind is None:
                continue
            left = ind.get(qn("w:left"))
            if left is not None and int(left) < 0:
                if ind.get(qn("w:hanging")):
                    ind.set(qn("w:left"), "180")
                else:
                    pPr.remove(ind)

        # docDefaults: single line spacing, no default after-gap. Word's
        # stock 1.15/8pt defaults are what made wrapped bullets look
        # double-spaced (Al's Jul 2026 style spec: single-spaced text,
        # gaps only where a style says so).
        styles_el = doc.styles.element
        pprd = styles_el.find(qn("w:docDefaults") + "/" + qn("w:pPrDefault"))
        if pprd is None:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            dd = styles_el.find(qn("w:docDefaults"))
            pprd = parse_xml(f'<w:pPrDefault {nsdecls("w")}/>')
            dd.append(pprd)
        ppr = pprd.find(qn("w:pPr"))
        if ppr is None:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            ppr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            pprd.append(ppr)
        sp = ppr.find(qn("w:spacing"))
        if sp is None:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            ppr.insert(0, sp)
        sp.set(qn("w:after"), "0")
        sp.set(qn("w:line"), "240")
        sp.set(qn("w:lineRule"), "auto")

        # 2) drop the VML horizontal-rule paragraphs
        for p in list(doc.paragraphs):
            if p._p.findall(".//" + qn("w:pict")):
                delete(p)

        # 3) section headings: bottom border + 6pt before
        for p in doc.paragraphs:
            if p.text.strip() in SECTION_HEADINGS:
                add_bottom_border(p)
                p.paragraph_format.space_before = Pt(6)

        # 4) bullet spacing -> style level (one knob), airy 6pt-after
        style = doc.styles["List Paragraph"]
        spf = style.paragraph_format
        spf.space_before = Pt(0)
        spf.space_after = Pt(6)
        st_pPr = style.element.find(qn("w:pPr"))
        if st_pPr is not None:
            cs = st_pPr.find(qn("w:contextualSpacing"))
            if cs is not None:  # would zero the gap between bullets
                st_pPr.remove(cs)
        for p in doc.paragraphs:
            if p.style.name.startswith("List"):
                pPr = p._p.pPr
                sp = pPr.find(qn("w:spacing")) if pPr is not None else None
                if sp is not None:
                    pPr.remove(sp)

        # 5) strip run-level font family overrides (TNR paste residue);
        #    the theme font is the single knob
        for r in doc.element.body.iter(qn("w:rFonts")):
            r.getparent().remove(r)

        # 6) content -> tags -------------------------------------------
        # Summary
        summary_p = find_one(doc, "Data scientist with")
        rebuild_runs(summary_p, [make_run(rpr(summary_p), "{{ summary }}")])

        # Jobs: USDA block is the prototype; the rest is deleted
        usda_h = find_one(doc, "USDA")
        usda_t = find_one(doc, "Data Scientist (Postgraduate")
        usda_bul = bullets_after(doc, usda_t)
        edu_head = find_one(doc, "Education")
        bold = rpr(usda_h)
        rebuild_runs(usda_h, [make_run(bold, "{{ job.org }}"),
                              make_run(bold, tab=True),
                              make_run(bold, "{{ job.dates }}")])
        ital = rpr(usda_t)
        rebuild_runs(usda_t, [make_run(ital, "{{ job.title }}"),
                              make_run(ital, tab=True),
                              make_run(ital, "{{ job.location }}")])
        proto = usda_bul[0]
        rebuild_runs(proto, [make_run(rpr(proto), "{{ b }}")])
        insert_before(usda_h, tag_paragraph("{%p for job in jobs %}"))
        insert_before(proto, tag_paragraph("{%p for b in job.bullets %}"))
        endfor_b = tag_paragraph("{%p endfor %}")
        insert_after(proto, endfor_b)
        endfor_b.addnext(tag_paragraph("{%p endfor %}"))
        for p in usda_bul[1:]:
            delete(p)
        for p in paragraphs_between(doc, usda_t, edu_head):
            if p._p.getparent() is not None and p._p is not proto._p:
                if p.text.strip().startswith(("{%p", "{{")):
                    continue
                delete(p)

        # Education: first line is the prototype
        proj_head = find_one(doc, "Project Experience")
        edu_lines = [p for p in paragraphs_between(doc, edu_head, proj_head)
                     if not p.text.strip().startswith("{%p")]
        e1 = edu_lines[0]
        e_bold = rpr(e1)
        e_ital = rpr(e1, lambda r: r.italic)
        if e_ital is None:
            e_ital = e_bold
        rebuild_runs(e1, [
            make_run(e_bold, "{{ e.degree }}"),
            make_run(e_ital,
                     "{% if e.institution %} – {{ e.institution }}{% endif %}"),
            make_run(e_bold, tab=True),
            make_run(e_bold, "{% if e.gpa %}{{ e.gpa }} GPA{% endif %}"),
            make_run(e_bold, tab=True),
            make_run(e_bold, "{{ e.dates }}"),
        ])
        insert_before(e1, tag_paragraph("{%p for e in education %}"))
        insert_after(e1, tag_paragraph("{%p endfor %}"))
        for p in edu_lines[1:]:
            delete(p)

        # Projects: Application Tracker block is the prototype
        app_h = find_one(doc, "Application Tracker")
        app_bul = bullets_after(doc, app_h)
        skills_head = find_one(doc, "Technical Skills")
        bold = rpr(app_h)
        rebuild_runs(app_h, [make_run(bold, "{{ p.org }}"),
                             make_run(bold, tab=True),
                             make_run(bold, "{{ p.dates }}")])
        proto = app_bul[0]
        rebuild_runs(proto, [make_run(rpr(proto), "{{ b }}")])
        insert_before(app_h, tag_paragraph("{%p for p in projects %}"))
        insert_before(proto, tag_paragraph("{%p for b in p.bullets %}"))
        endfor_b = tag_paragraph("{%p endfor %}")
        insert_after(proto, endfor_b)
        endfor_b.addnext(tag_paragraph("{%p endfor %}"))
        for p in app_bul[1:]:
            delete(p)
        for p in paragraphs_between(doc, app_h, skills_head):
            if p._p.getparent() is not None and p._p is not proto._p:
                if p.text.strip().startswith("{%p"):
                    continue
                delete(p)

        # Skills: first labeled bullet is the prototype
        pubs_head = find_one(doc, "Publications and Awards")
        skill_ps = [p for p in paragraphs_between(doc, skills_head, pubs_head)
                    if not p.text.strip().startswith("{%p")]
        s1 = skill_ps[0]
        s_bold = rpr(s1)
        s_plain = rpr(s1, lambda r: not r.bold)
        if s_plain is None:
            s_plain = s_bold
        rebuild_runs(s1, [make_run(s_bold, "{{ s.label }}: "),
                          make_run(s_plain, "{{ s.text }}")])
        insert_before(s1, tag_paragraph("{%p for s in skills %}"))
        insert_after(s1, tag_paragraph("{%p endfor %}"))
        for p in skill_ps[1:]:
            delete(p)

        # Publications: first bullet is the prototype (v1 renders the
        # whole citation flat — the bold right-tabbed date column needs a
        # master grammar for dates first; noted for M3)
        pub_ps = bullets_after(doc, pubs_head)
        p1 = pub_ps[0]
        p_plain = rpr(p1, lambda r: not r.bold)
        if p_plain is None:
            p_plain = rpr(p1)
        rebuild_runs(p1, [make_run(p_plain, "{{ pub }}")])
        insert_before(p1, tag_paragraph("{%p for pub in publications %}"))
        insert_after(p1, tag_paragraph("{%p endfor %}"))
        for p in pub_ps[1:]:
            delete(p)

        doc.save(str(dst))
        print(f"wrote {dst}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit(__doc__)
    hollow(Path(sys.argv[1]), Path(sys.argv[2]))
