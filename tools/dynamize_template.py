"""Build a dynamic-header template: the name and contact lines render
from the context (header.name / header.contact_lines) instead of being
static text — required for loose mode / other users. Re-runnable.

Usage: python tools/dynamize_template.py TEMPLATE.docx OUT.docx
"""

import copy
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _set_text(p, s):
    rpr = None
    for r in p._p.findall(qn("w:r")):
        rr = r.find(qn("w:rPr"))
        if rr is not None:
            rpr = copy.deepcopy(rr)
            break
    if rpr is None:
        for h in p._p.findall(qn("w:hyperlink")):
            for r in h.findall(qn("w:r")):
                rr = r.find(qn("w:rPr"))
                if rr is not None:
                    rpr = copy.deepcopy(rr)
                    st = rpr.find(qn("w:rStyle"))
                    if st is not None:
                        rpr.remove(st)
                    break
            if rpr is not None:
                break
    for child in list(p._p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            p._p.remove(child)
    nr = OxmlElement("w:r")
    if rpr is not None:
        nr.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = s
    nr.append(t)
    p._p.append(nr)


def _tag_para(text):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def main(src, dst):
    doc = Document(src)
    name_p = next(p for p in doc.paragraphs if p.text.strip())
    contact = []
    started = False
    for p in doc.paragraphs:
        if p._p is name_p._p:
            started = True
            continue
        if not started:
            continue
        pPr = p._p.pPr
        if pPr is not None and pPr.find(qn("w:pBdr")) is not None:
            break
        if p.text.strip() or p._p.findall(qn("w:hyperlink")):
            contact.append(p)
    if not contact:
        raise SystemExit("no contact paragraph found under the name line")
    proto = contact[0]
    for extra in contact[1:]:
        extra._p.getparent().remove(extra._p)
    _set_text(name_p, "{{ header.name }}")
    _set_text(proto, "{{ hl }}")
    proto._p.addprevious(_tag_para("{%p for hl in header.contact_lines %}"))
    proto._p.addnext(_tag_para("{%p endfor %}"))
    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit(__doc__)
    main(sys.argv[1], sys.argv[2])
