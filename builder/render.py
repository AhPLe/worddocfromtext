"""M2 render: render context -> .docx via docxtpl.

docxtpl is NOT installable in Claude's sandbox (no network), so
render_docx is parse-checked there and execution-verified on the
owner's machine (pip install docxtpl). normalize_right_tabs uses only
python-docx and IS execution-verified in the sandbox.

The error gate applies upstream (cli.py): rendering is refused while
any ERROR finding exists; slot markers can never reach this module
because context.rescan_for_slots runs at build time.
"""

import re
from pathlib import Path


def render_docx(context, template_path, out_path, spacing="template"):
    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required for rendering — install it locally with: "
            "pip install docxtpl") from exc
    tpl = DocxTemplate(str(template_path))
    # autoescape=True: content contains literal '&' (Novelis 'R&T',
    # 'Data & technical'); unescaped it corrupts the document XML.
    # Parameter name per docxtpl docs — the Novelis line rendering
    # correctly is the first-run canary (see README verification).
    tpl.render(context, autoescape=True)
    tpl.save(str(out_path))
    normalize_right_tabs(out_path)
    if spacing != "template":
        apply_spacing(out_path, spacing)
    linkify_header(out_path)


def normalize_right_tabs(path):
    """Glue right tab stops to the right margin.

    The date/GPA columns are right-anchored tabs. If the template's
    margins change (a one-edit knob in Word), every right tab shifts by
    the same delta so the columns follow the new text width — no manual
    tab-stop bookkeeping.
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    sec = doc.sections[0]
    width = (sec.page_width.twips - sec.left_margin.twips
             - sec.right_margin.twips)

    right_tabs = []
    for p in doc.paragraphs:
        pPr = p._p.pPr
        tabs = pPr.find(qn("w:tabs")) if pPr is not None else None
        if tabs is None:
            continue
        for t in tabs.findall(qn("w:tab")):
            if t.get(qn("w:val")) == "right":
                right_tabs.append(t)
    if not right_tabs:
        return
    rightmost = max(int(t.get(qn("w:pos"))) for t in right_tabs)
    delta = width - rightmost
    if delta == 0:
        return
    for t in right_tabs:
        t.set(qn("w:pos"), str(int(t.get(qn("w:pos"))) + delta))
    doc.save(str(path))


def apply_spacing(path, mode):
    """Toggle the bullet-spacing knob on a rendered (or template) docx.

    'airy'  = 6pt after every bullet (List Paragraph style spacing).
    'tight' = same values plus contextualSpacing: Word suppresses the
              gap between consecutive same-style paragraphs, so bullets
              in a group sit at single spacing and only the last keeps
              its 6pt before the next heading. Worth roughly 1.6" of
              vertical space on a full render — the difference between
              one page and two.
    """
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Pt

    doc = Document(str(path))
    # docDefaults: single line spacing, no default after-gap — Word's
    # stock 1.15/8pt defaults pad every non-bullet paragraph too, which
    # is what pushed otherwise-tight renders past one page.
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    pprd = dd.find(qn("w:pPrDefault")) if dd is not None else None
    if pprd is None and dd is not None:
        pprd = parse_xml(f'<w:pPrDefault {nsdecls("w")}/>')
        dd.append(pprd)
    if pprd is not None:
        ppr = pprd.find(qn("w:pPr"))
        if ppr is None:
            ppr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            pprd.append(ppr)
        sp = ppr.find(qn("w:spacing"))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            ppr.insert(0, sp)
        sp.set(qn("w:after"), "0")
        sp.set(qn("w:line"), "240")
        sp.set(qn("w:lineRule"), "auto")
    style = doc.styles["List Paragraph"]
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0  # single-spaced within a wrapped bullet
    pPr = style.element.find(qn("w:pPr"))
    cs = pPr.find(qn("w:contextualSpacing"))
    if mode == "tight":
        if cs is None:
            el = parse_xml(f'<w:contextualSpacing {nsdecls("w")}/>')
            anchor = pPr.find(qn("w:ind"))
            if anchor is None:
                anchor = pPr.find(qn("w:spacing"))
            if anchor is not None:
                anchor.addnext(el)
            else:
                pPr.append(el)
    elif mode == "airy":
        if cs is not None:
            pPr.remove(cs)
    else:
        raise ValueError(f"unknown spacing mode: {mode!r}")
    doc.save(str(path))


_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+")
_URL_RE = re.compile(
    r"(?:https?://\S+|(?:www\.)?[\w-]+\.github\.io\S*|"
    r"(?:www\.)?linkedin\.com/\S+)", re.I)


def linkify_header(path):
    """Turn plain email / GitHub / LinkedIn text in the header zone into
    real hyperlinks (Hyperlink style, external rels). The zone is the
    paragraphs before the first bordered section heading, capped at 6 —
    so DOIs in Publications are never touched. Runs already inside
    hyperlinks are untouched (python-docx doesn't surface them), so the
    static-header template is safe too."""
    import copy as _copy
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(path))
    seen = 0
    for p in doc.paragraphs[:12]:
        pPr = p._p.pPr
        if pPr is not None and pPr.find(qn("w:pBdr")) is not None:
            break
        seen += 1
        if seen > 6:
            break
        for run in list(p.runs):
            text = run.text
            if not text:
                continue
            spans = [(m.start(), m.end(), "mailto:" + m.group(0))
                     for m in _EMAIL_RE.finditer(text)]
            for m in _URL_RE.finditer(text):
                if any(s <= m.start() < e for s, e, _ in spans):
                    continue
                url = m.group(0).rstrip(".,;")
                tgt = url if url.lower().startswith("http")                     else "https://" + url
                spans.append((m.start(), m.start() + len(url), tgt))
            if not spans:
                continue
            spans.sort()
            rpr = run._r.find(qn("w:rPr"))
            parent = run._r.getparent()
            idx = parent.index(run._r)
            parent.remove(run._r)

            def mkrun(txt, link=False):
                nr = OxmlElement("w:r")
                if rpr is not None:
                    nr.append(_copy.deepcopy(rpr))
                if link:
                    rp = nr.find(qn("w:rPr"))
                    if rp is None:
                        rp = OxmlElement("w:rPr")
                        nr.insert(0, rp)
                    st = OxmlElement("w:rStyle")
                    st.set(qn("w:val"), "Hyperlink")
                    rp.insert(0, st)
                wt = OxmlElement("w:t")
                wt.set(qn("xml:space"), "preserve")
                wt.text = txt
                nr.append(wt)
                return nr

            pos = 0
            for s, e, tgt in spans:
                if s > pos:
                    parent.insert(idx, mkrun(text[pos:s]))
                    idx += 1
                rid = doc.part.relate_to(
                    tgt, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                h = OxmlElement("w:hyperlink")
                h.set(qn("r:id"), rid)
                h.append(mkrun(text[s:e], link=True))
                parent.insert(idx, h)
                idx += 1
                pos = e
            if pos < len(text):
                parent.insert(idx, mkrun(text[pos:]))
                idx += 1
    doc.save(str(path))
